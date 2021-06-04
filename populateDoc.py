from docx import Document
from docx.shared import Inches
import random
import os.path
from os import path

document = Document()
randomlist = random.sample(range(0, 2119), 2119)


for i in randomlist:
	p = document.add_paragraph()
	r = p.add_run()
	r.add_break()
	if path.exists(str(i) + '.jpg'):
		r.add_picture(str(i) + '.jpg')
	elif path.exists(str(i) + '.png'):
		r.add_picture(str(i) + '.png')
	else:
		continue
	r.add_text(str(i))


document.save('book.docx')