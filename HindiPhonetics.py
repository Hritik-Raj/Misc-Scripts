
import docx
# import epitran
import transliterator

from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement

def func(value):
    return ''.join(value.splitlines())

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para
# epi = epitran.Epitran('hin-Deva')
start = False
count = 0
doc = docx.Document('/Users/hritikraj/Downloads/Hindi.docx')
hindistr = None
startlines = ["Translation in Devanagari:", "Translation in Devanagari", " Translation in Devanagari:"]
hindiarray = []
for i in doc.paragraphs:
    # print(type(i.text))
    # print(i.text)
    if (i.text == "Translation in Devanagari:" or i.text == "Translation in Devanagari" or i.text == " Translation in Devanagari:"):
        hindistr = ""
        start = True
    if (i.text == "Transliteration in Roman Script:"):
        hindiarray.append(hindistr)
        start = False
        print("-------")
    if start == True and i.text not in startlines:
        hindistr += i.text
        # transliteration = epi.transliterate(hindistr)
        
        # hindistr = func(hindistr)
        # print(hindistr)
        # translit = transliterator.transliterate(hindistr, 'devanagari', 'harvardkyoto')
        # print(translit)
        # hindistr = ""

# for item in hindiarray:
#     item = transliterator.transliterate(item, 'devanagari', 'harvardkyoto',  {'outputASCIIEncoded' : True})
    # print(translit)
    # print(item)
    # print("===============================================")


for i in doc.paragraphs:
    if (i.text == "Transliteration in Roman Script:"):
        x = hindiarray[count]
        x = transliterator.transliterate(x, 'devanagari', 'harvardkyoto')
        # x = ''.join(c for c in x if valid_xml_char_ordinal(c))
        insert_paragraph_after(i, x.decode('utf-8'))
        count += 1
        

print(count)

doc.save('/Users/hritikraj/Downloads/Hindi.docx')










    # if (i.text == "Translation in Devanagari:"):
    #     start = True
    #     # while (i.text != "Transliteration in Roman Script:"):
                      